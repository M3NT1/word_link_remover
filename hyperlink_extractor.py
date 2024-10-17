# hyperlink_extractor.py

import logging
from urllib.parse import urlparse
import re
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.shared import qn
from collections import defaultdict

logging.basicConfig(filename='hyperlink_extractor.log', level=logging.DEBUG,
                    format='%(asctime)s - %(levelname)s - %(message)s')


def extract_hyperlinks(document):
    hyperlinks = []
    print("Kezdődik a hivatkozások kinyerése")
    logging.info("Kezdődik a hivatkozások kinyerése")

    # 1. Beágyazott hiperhivatkozások
    print("1. módszer: Beágyazott hiperhivatkozások keresése")
    logging.info("1. módszer: Beágyazott hiperhivatkozások keresése")
    for i, paragraph in enumerate(document.paragraphs):
        if i % 100 == 0:
            print(f"  Feldolgozás alatt: {i}. bekezdés")
        for run in paragraph.runs:
            try:
                if hasattr(run.element, 'hyperlink'):
                    if run.element.hyperlink is not None:
                        target = run.element.hyperlink.get(qn('r:id'))
                        if target:
                            target = document.part.rels[target].target_ref
                            context = get_context(paragraph, run)
                            hyperlinks.append({
                                "text": run.text,
                                "target": target,
                                "type": determine_hyperlink_type(target),
                                "context": context,
                                "link_text": run.text
                            })
                            print(f"    Talált beágyazott hivatkozás: {run.text} -> {target}")
                            logging.debug(f"Talált beágyazott hivatkozás: {run.text} -> {target}")
                else:
                    for elem in run._element.xpath('.//w:hyperlink'):
                        rId = elem.get(qn('r:id'))
                        if rId:
                            target = document.part.rels[rId].target_ref
                            context = get_context(paragraph, run)
                            hyperlinks.append({
                                "text": run.text,
                                "target": target,
                                "type": determine_hyperlink_type(target),
                                "context": context,
                                "link_text": run.text
                            })
                            print(f"    Talált beágyazott hivatkozás (alternatív módszer): {run.text} -> {target}")
                            logging.debug(f"Talált beágyazott hivatkozás (alternatív módszer): {run.text} -> {target}")
            except AttributeError as e:
                logging.warning(f"AttributeError a hivatkozás kinyerése során: {str(e)}")

    # 2. Szöveges URL-ek
    print("2. módszer: Szöveges URL-ek keresése")
    logging.info("2. módszer: Szöveges URL-ek keresése")
    url_pattern = re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+')
    for i, paragraph in enumerate(document.paragraphs):
        if i % 100 == 0:
            print(f"  Feldolgozás alatt: {i}. bekezdés")
        urls = url_pattern.findall(paragraph.text)
        for url in urls:
            context = get_context(paragraph, url)
            hyperlinks.append({
                "text": url,
                "target": url,
                "type": determine_hyperlink_type(url),
                "context": context,
                "link_text": url
            })
            print(f"    Talált szöveges URL: {url}")
            logging.debug(f"Talált szöveges URL: {url}")

    # 3. Mezők (fields)
    print("3. módszer: Mezők keresése")
    logging.info("3. módszer: Mezők keresése")
    for i, paragraph in enumerate(document.paragraphs):
        if i % 100 == 0:
            print(f"  Feldolgozás alatt: {i}. bekezdés")
        for run in paragraph.runs:
            if run._element.get('fldChar') is not None:
                field_text = get_field_text(run._element)
                if field_text.startswith('HYPERLINK'):
                    parts = field_text.split('"')
                    if len(parts) > 1:
                        url = parts[1]
                        context = get_context(paragraph, run)
                        hyperlinks.append({
                            "text": "Mező hivatkozás",
                            "target": url,
                            "type": determine_hyperlink_type(url),
                            "context": context,
                            "link_text": run.text
                        })
                        print(f"    Talált mező hivatkozás: {url}")
                        logging.debug(f"Talált mező hivatkozás: {url}")

    # 4. Könyvjelzők
    print("4. módszer: Könyvjelzők keresése")
    logging.info("4. módszer: Könyvjelzők keresése")
    for i, bookmark in enumerate(document.element.xpath('//w:bookmarkStart')):
        if i % 100 == 0:
            print(f"  Feldolgozás alatt: {i}. könyvjelző")
        name = bookmark.get(qn('w:name'))
        if name.startswith('_'):
            continue  # Belső könyvjelzők kihagyása
        # Keressük meg a könyvjelzőhöz tartozó szöveget
        bookmark_text = get_bookmark_text(document, name)
        hyperlinks.append({
            "text": f"Könyvjelző: {name}",
            "target": f"#{name}",
            "type": "belső",
            "context": "Könyvjelző",
            "link_text": bookmark_text
        })
        print(f"    Talált könyvjelző: {name}")
        logging.debug(f"Talált könyvjelző: {name}")

    # 5. Speciális belső hivatkozások
    print("5. módszer: Speciális belső hivatkozások keresése")
    logging.info("5. módszer: Speciális belső hivatkozások keresése")
    internal_link_pattern = re.compile(r'(.*?)\t.*?(Dokumentum belső hivatkozás: BKM_[A-F0-9_]+)')
    for i, paragraph in enumerate(document.paragraphs):
        if i % 100 == 0:
            print(f"  Feldolgozás alatt: {i}. bekezdés")
        matches = internal_link_pattern.findall(paragraph.text)
        for match in matches:
            text, target = match
            text = text.strip()
            target = target.strip()
            link_type = determine_internal_link_type(target)
            context = get_context(paragraph, text)
            hyperlinks.append({
                "text": text,
                "target": target,
                "type": link_type,
                "context": context,
                "link_text": text
            })
            print(f"    Talált speciális belső hivatkozás: {text} -> {target} ({link_type})")
            logging.debug(f"Talált speciális belső hivatkozás: {text} -> {target} ({link_type})")

    unique_links = remove_duplicates(hyperlinks)
    print(f"Összesen {len(unique_links)} egyedi hivatkozás található")
    logging.info(f"Összesen {len(unique_links)} egyedi hivatkozás található")

    # Hivatkozások csoportosítása és statisztika
    grouped_links = group_links(unique_links)
    print_link_statistics(grouped_links)

    return unique_links


def get_context(paragraph, text):
    full_text = paragraph.text
    if isinstance(text, str):
        start_index = full_text.index(text)
        end_index = start_index + len(text)
    else:  # Ha a text egy run objektum
        start_index = full_text.index(text.text)
        end_index = start_index + len(text.text)

    # Kontextus: 50 karakter a hivatkozás előtt és után
    context_start = max(0, start_index - 50)
    context_end = min(len(full_text), end_index + 50)

    return full_text[context_start:context_end]


def get_field_text(element):
    field_text = ''
    for e in element.getparent().iter():
        if e.tag.endswith('fldChar') and e.get('fldCharType') == 'end':
            break
        if e.tag.endswith('instrText'):
            field_text += e.text
    return field_text


def determine_hyperlink_type(url):
    if url.startswith("#") or "BKM_" in url:
        return "belső"
    elif is_valid_url(url):
        return "külső"
    else:
        return "törött"


def determine_internal_link_type(target):
    if "Dokumentum belső hivatkozás: BKM_" in target:
        if "részleges egyezés" in target.lower():
            return "Érvényes belső hivatkozás (részleges egyezés)"
        else:
            return "Valószínűleg törött belső hivatkozás (szellem-hivatkozás)"
    return "Ismeretlen belső hivatkozás típus"


def is_valid_url(url):
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc])
    except:
        return False


def remove_duplicates(hyperlinks):
    seen = set()
    unique_hyperlinks = []
    for link in hyperlinks:
        key = (link['text'], link['target'])
        if key not in seen:
            seen.add(key)
            unique_hyperlinks.append(link)
    return unique_hyperlinks


def group_links(hyperlinks):
    grouped = defaultdict(list)
    for link in hyperlinks:
        grouped[link['type']].append(link)
    return grouped


def print_link_statistics(grouped_links):
    print("\nHivatkozás statisztika:")
    for link_type, links in grouped_links.items():
        print(f"{link_type}: {len(links)} db")
    print(f"Összes hivatkozás: {sum(len(links) for links in grouped_links.values())} db")


def get_bookmark_text(document, bookmark_name):
    for paragraph in document.paragraphs:
        if bookmark_name in paragraph.text:
            return paragraph.text
    return ""
